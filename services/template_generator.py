import os
import json
import hashlib
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, asdict
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import tempfile

# Import storage service
from services.storage import storage_service, StorageError

logger = logging.getLogger(__name__)

@dataclass
class ContentSection:
    """Represents a section in the proposal template"""
    heading: str
    instruction: str
    placeholder: str

@dataclass
class ContentModel:
    """Deterministic content model for proposal templates"""
    cover: Dict[str, str]
    sections: List[ContentSection]
    funder_notes: Optional[str]
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for hashing"""
        return {
            'cover': self.cover,
            'sections': [asdict(section) for section in self.sections],
            'funder_notes': self.funder_notes,
            'metadata': self.metadata
        }
    
    def compute_hash(self) -> str:
        """Compute stable SHA256 hash from sorted content"""
        # Convert to sorted JSON for deterministic hashing
        content_dict = self.to_dict()
        content_json = json.dumps(content_dict, sort_keys=True, separators=(',', ':'))
        return hashlib.sha256(content_json.encode('utf-8')).hexdigest()

class TemplateBuildError(Exception):
    """Exception for template generation failures"""
    pass

class PDFGenerationError(Exception):
    """Exception for PDF generation failures"""
    pass

class ProposalTemplateGenerator:
    """Enhanced proposal template generator with deterministic output and PDF support"""
    
    def __init__(self):
        self.pdf_engine = os.getenv("PDF_ENGINE", "reportlab").lower()
        self._check_pdf_capabilities()
    
    def _check_pdf_capabilities(self):
        """Check available PDF generation capabilities"""
        if self.pdf_engine == "weasyprint":
            try:
                import weasyprint
                logger.info("✅ WeasyPrint PDF engine available")
            except ImportError:
                logger.warning("⚠️ WeasyPrint not available, falling back to ReportLab")
                self.pdf_engine = "reportlab"
        
        if self.pdf_engine == "reportlab":
            try:
                import reportlab
                logger.info("✅ ReportLab PDF engine available")
            except ImportError:
                logger.error("❌ No PDF engine available - PDF generation disabled")
                self.pdf_engine = None
    
    def build_content_model(
        self, 
        opportunity_data: Dict[str, Any], 
        sections: List[Dict[str, str]], 
        funder_notes: Optional[str] = None,
        hints: Optional[Dict[str, str]] = None
    ) -> ContentModel:
        """Build deterministic content model from opportunity data and sections"""
        try:
            # Extract opportunity information
            title = opportunity_data.get('title', 'Funding Opportunity')
            donor = opportunity_data.get('donor', 'Not specified')
            deadline = opportunity_data.get('deadline', 'Not specified')
            amount = opportunity_data.get('amount', 'Not specified')
            location = opportunity_data.get('location', 'Not specified')
            themes = opportunity_data.get('themes', [])
            opportunity_url = opportunity_data.get('opportunity_url', 'Not provided')
            
            # Apply hints if provided
            if hints:
                org_name = hints.get('org_name', 'Your Organization')
                country = hints.get('country', 'Your Country')
                contact_name = hints.get('contact_name', 'Your Name')
            else:
                org_name = 'Your Organization'
                country = 'Your Country'
                contact_name = 'Your Name'
            
            # Build cover section
            cover = {
                'title': title,
                'donor': donor,
                'deadline': deadline,
                'amount': amount,
                'location': location,
                'themes': ', '.join(themes) if isinstance(themes, list) else str(themes),
                'opportunity_url': opportunity_url,
                'org_name': org_name,
                'country': country,
                'contact_name': contact_name
            }
            
            # Build sections with placeholders
            content_sections = []
            for section in sections:
                heading = section.get('heading', 'Section')
                instruction = section.get('instruction', 'No instruction provided')
                
                # Create placeholder text
                placeholder = f"[Your content for {heading.lower()} goes here. {instruction}]"
                
                content_sections.append(ContentSection(
                    heading=heading,
                    instruction=instruction,
                    placeholder=placeholder
                ))
            
            # Build metadata
            metadata = {
                'generated_at': datetime.now().isoformat(),
                'opportunity_id': opportunity_data.get('id'),
                'source_url': opportunity_data.get('source_url'),
                'version': '2.0.0'
            }
            
            return ContentModel(
                cover=cover,
                sections=content_sections,
                funder_notes=funder_notes,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ Failed to build content model: {e}")
            raise TemplateBuildError(f"Content model build failed: {e}")
    
    def generate_docx(self, content_model: ContentModel) -> bytes:
        """Generate DOCX document from content model"""
        try:
            doc = Document()
            
            # Add title page
            title_para = doc.add_heading(f"Proposal Template: {content_model.cover['title']}", 0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add cover information table
            doc.add_heading('Opportunity Information', level=1)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            cover_fields = [
                ('Donor/Funder', content_model.cover['donor']),
                ('Deadline', content_model.cover['deadline']),
                ('Funding Amount', content_model.cover['amount']),
                ('Location/Eligibility', content_model.cover['location']),
                ('Themes/Focus Areas', content_model.cover['themes']),
                ('Opportunity URL', content_model.cover['opportunity_url']),
                ('Organization', content_model.cover['org_name']),
                ('Country', content_model.cover['country']),
                ('Contact Person', content_model.cover['contact_name'])
            ]
            
            for label, value in cover_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].text = str(value)
            
            # Add funder notes if provided
            if content_model.funder_notes:
                doc.add_paragraph('')
                doc.add_heading('Funder Requirements & Notes', level=1)
                funder_para = doc.add_paragraph(content_model.funder_notes.strip())
                funder_para.style = 'Intense Quote'
            
            # Add instructions
            doc.add_paragraph('')
            instruction_para = doc.add_paragraph()
            instruction_para.add_run('Instructions: ').bold = True
            instruction_para.add_run('This template provides the structure for your proposal. Replace the instructional text below with your actual content. Each section includes guidance on what to include.')
            
            # Add proposal sections
            doc.add_paragraph('')
            doc.add_heading('Proposal Sections', level=1)
            
            for i, section in enumerate(content_model.sections, 1):
                # Section heading
                doc.add_heading(f"{i}. {section.heading}", level=2)
                
                # Instruction
                instruction_para = doc.add_paragraph()
                instruction_para.add_run('[INSTRUCTION] ').bold = True
                instruction_para.add_run(section.instruction)
                instruction_para.style = 'Subtle Emphasis'
                
                # Placeholder
                content_para = doc.add_paragraph()
                content_para.add_run(section.placeholder)
                content_para.style = 'Quote'
                
                doc.add_paragraph('')
            
            # Add footer
            doc.add_page_break()
            footer_heading = doc.add_heading('Template Information', level=2)
            footer_para = doc.add_paragraph()
            footer_para.add_run('Generated: ').bold = True
            footer_para.add_run(content_model.metadata['generated_at'])
            footer_para.add_run('\nBy: ').bold = True
            footer_para.add_run('NGOInfo ReqAgent')
            footer_para.add_run('\nVersion: ').bold = True
            footer_para.add_run(content_model.metadata['version'])
            
            # Save to bytes
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                doc.save(tmp_file.name)
                with open(tmp_file.name, 'rb') as f:
                    docx_bytes = f.read()
                os.unlink(tmp_file.name)
            
            logger.info(f"✅ Generated DOCX: {len(docx_bytes)} bytes")
            return docx_bytes
            
        except Exception as e:
            logger.error(f"❌ DOCX generation failed: {e}")
            raise TemplateBuildError(f"DOCX generation failed: {e}")
    
    def generate_pdf(self, content_model: ContentModel) -> Optional[bytes]:
        """Generate PDF from content model using available engine"""
        if not self.pdf_engine:
            logger.warning("⚠️ PDF generation disabled - no engine available")
            return None
        
        try:
            if self.pdf_engine == "weasyprint":
                return self._generate_pdf_weasyprint(content_model)
            elif self.pdf_engine == "reportlab":
                return self._generate_pdf_reportlab(content_model)
            else:
                logger.error(f"❌ Unknown PDF engine: {self.pdf_engine}")
                return None
                
        except Exception as e:
            logger.error(f"❌ PDF generation failed: {e}")
            raise PDFGenerationError(f"PDF generation failed: {e}")
    
    def _generate_pdf_weasyprint(self, content_model: ContentModel) -> bytes:
        """Generate PDF using WeasyPrint"""
        try:
            from weasyprint import HTML, CSS
            from jinja2 import Template
            
            # HTML template
            html_template = """
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
                    h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
                    h2 { color: #34495e; margin-top: 30px; }
                    .cover-table { width: 100%; border-collapse: collapse; margin: 20px 0; }
                    .cover-table td { padding: 8px; border: 1px solid #ddd; }
                    .cover-table td:first-child { font-weight: bold; background: #f8f9fa; }
                    .section { margin: 20px 0; }
                    .instruction { background: #f8f9fa; padding: 10px; border-left: 4px solid #3498db; }
                    .placeholder { background: #fff3cd; padding: 10px; border: 1px solid #ffeaa7; }
                    .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.9em; color: #666; }
                </style>
            </head>
            <body>
                <h1>Proposal Template: {{ cover.title }}</h1>
                
                <h2>Opportunity Information</h2>
                <table class="cover-table">
                    <tr><td>Donor/Funder</td><td>{{ cover.donor }}</td></tr>
                    <tr><td>Deadline</td><td>{{ cover.deadline }}</td></tr>
                    <tr><td>Funding Amount</td><td>{{ cover.amount }}</td></tr>
                    <tr><td>Location/Eligibility</td><td>{{ cover.location }}</td></tr>
                    <tr><td>Themes/Focus Areas</td><td>{{ cover.themes }}</td></tr>
                    <tr><td>Opportunity URL</td><td>{{ cover.opportunity_url }}</td></tr>
                    <tr><td>Organization</td><td>{{ cover.org_name }}</td></tr>
                    <tr><td>Country</td><td>{{ cover.country }}</td></tr>
                    <tr><td>Contact Person</td><td>{{ cover.contact_name }}</td></tr>
                </table>
                
                {% if funder_notes %}
                <h2>Funder Requirements & Notes</h2>
                <div class="instruction">{{ funder_notes }}</div>
                {% endif %}
                
                <h2>Instructions</h2>
                <p>This template provides the structure for your proposal. Replace the instructional text below with your actual content. Each section includes guidance on what to include.</p>
                
                <h2>Proposal Sections</h2>
                {% for section in sections %}
                <div class="section">
                    <h3>{{ loop.index }}. {{ section.heading }}</h3>
                    <div class="instruction"><strong>[INSTRUCTION]</strong> {{ section.instruction }}</div>
                    <div class="placeholder">{{ section.placeholder }}</div>
                </div>
                {% endfor %}
                
                <div class="footer">
                    <p><strong>Generated:</strong> {{ metadata.generated_at }}</p>
                    <p><strong>By:</strong> NGOInfo ReqAgent</p>
                    <p><strong>Version:</strong> {{ metadata.version }}</p>
                </div>
            </body>
            </html>
            """
            
            # Render HTML
            template = Template(html_template)
            html_content = template.render(
                cover=content_model.cover,
                sections=content_model.sections,
                funder_notes=content_model.funder_notes,
                metadata=content_model.metadata
            )
            
            # Generate PDF
            html = HTML(string=html_content)
            pdf_bytes = html.write_pdf()
            
            logger.info(f"✅ Generated PDF with WeasyPrint: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except ImportError:
            logger.error("❌ WeasyPrint not available")
            raise PDFGenerationError("WeasyPrint not available")
    
    def _generate_pdf_reportlab(self, content_model: ContentModel) -> bytes:
        """Generate PDF using ReportLab (pure Python fallback)"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from io import BytesIO
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12
            )
            
            # Build story
            story = []
            
            # Title
            story.append(Paragraph(f"Proposal Template: {content_model.cover['title']}", title_style))
            story.append(Spacer(1, 20))
            
            # Cover information table
            story.append(Paragraph("Opportunity Information", heading_style))
            
            cover_data = [
                ['Donor/Funder', content_model.cover['donor']],
                ['Deadline', content_model.cover['deadline']],
                ['Funding Amount', content_model.cover['amount']],
                ['Location/Eligibility', content_model.cover['location']],
                ['Themes/Focus Areas', content_model.cover['themes']],
                ['Opportunity URL', content_model.cover['opportunity_url']],
                ['Organization', content_model.cover['org_name']],
                ['Country', content_model.cover['country']],
                ['Contact Person', content_model.cover['contact_name']]
            ]
            
            cover_table = Table(cover_data, colWidths=[2*inch, 4*inch])
            cover_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(cover_table)
            story.append(Spacer(1, 20))
            
            # Funder notes
            if content_model.funder_notes:
                story.append(Paragraph("Funder Requirements & Notes", heading_style))
                story.append(Paragraph(content_model.funder_notes, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Instructions
            story.append(Paragraph("Instructions", heading_style))
            story.append(Paragraph("This template provides the structure for your proposal. Replace the instructional text below with your actual content. Each section includes guidance on what to include.", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Sections
            story.append(Paragraph("Proposal Sections", heading_style))
            
            for i, section in enumerate(content_model.sections, 1):
                story.append(Paragraph(f"{i}. {section.heading}", heading_style))
                story.append(Paragraph(f"[INSTRUCTION] {section.instruction}", styles['Normal']))
                story.append(Paragraph(section.placeholder, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph("Template Information", heading_style))
            story.append(Paragraph(f"Generated: {content_model.metadata['generated_at']}", styles['Normal']))
            story.append(Paragraph("By: NGOInfo ReqAgent", styles['Normal']))
            story.append(Paragraph(f"Version: {content_model.metadata['version']}", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info(f"✅ Generated PDF with ReportLab: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except ImportError:
            logger.error("❌ ReportLab not available")
            raise PDFGenerationError("ReportLab not available")
    
    def generate_template(
        self, 
        opportunity_data: Dict[str, Any], 
        sections: List[Dict[str, str]], 
        funder_notes: Optional[str] = None,
        hints: Optional[Dict[str, str]] = None
    ) -> Tuple[ContentModel, bytes, Optional[bytes]]:
        """Generate complete template with content model, DOCX, and PDF"""
        try:
            # Build content model
            content_model = self.build_content_model(opportunity_data, sections, funder_notes, hints)
            
            # Generate DOCX
            docx_bytes = self.generate_docx(content_model)
            
            # Generate PDF (optional)
            pdf_bytes = None
            try:
                pdf_bytes = self.generate_pdf(content_model)
            except PDFGenerationError as e:
                logger.warning(f"⚠️ PDF generation failed, continuing with DOCX only: {e}")
            
            return content_model, docx_bytes, pdf_bytes
            
        except Exception as e:
            logger.error(f"❌ Template generation failed: {e}")
            raise TemplateBuildError(f"Template generation failed: {e}")

