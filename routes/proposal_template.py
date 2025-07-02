from fastapi import APIRouter, HTTPException, Depends, status, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
import logging
import os
from datetime import datetime
from typing import Dict, Any, Optional
import uuid
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import traceback
import shutil
import tempfile

# Database imports
from db import get_db
from models import FundingOpportunity, StatusEnum
from schemas import CreateProposalTemplateRequest, ProposalTemplateResponse
from utils.auth import require_admin_auth

# Set up logging
logger = logging.getLogger(__name__)

# Create router
router = APIRouter(prefix="/admin", tags=["proposal-template"])

# Set up Jinja2 templates
templates = Jinja2Templates(directory="templates")

class ProposalTemplateGenerator:
    """Service for generating proposal template documents"""
    
    def __init__(self):
        # Use system temp directory for generated templates (cross-platform, always writable)
        self.templates_dir = os.path.join(tempfile.gettempdir(), "reqagent_templates")
        self._ensure_templates_directory()
    
    def _ensure_templates_directory(self):
        """Ensure templates directory exists and has proper permissions"""
        try:
            # Create directory if it doesn't exist
            if not os.path.exists(self.templates_dir):
                logger.info(f"📁 Creating templates directory: {self.templates_dir}")
                os.makedirs(self.templates_dir, exist_ok=True)
            
            # Check directory permissions
            if not os.access(self.templates_dir, os.W_OK):
                logger.error(f"🔴 Templates directory is not writable: {self.templates_dir}")
                raise PermissionError(f"Templates directory is not writable: {self.templates_dir}")
            
            # Check disk space (require at least 100MB)
            disk_usage = shutil.disk_usage(self.templates_dir)
            free_space_mb = disk_usage.free / (1024 * 1024)
            
            if free_space_mb < 100:
                logger.warning(f"⚠️ Low disk space: {free_space_mb:.1f}MB available")
                if free_space_mb < 10:
                    raise OSError(f"Insufficient disk space: {free_space_mb:.1f}MB available (minimum 10MB required)")
            
            logger.info(f"✅ Templates directory ready: {self.templates_dir} ({free_space_mb:.1f}MB available)")
            
        except Exception as e:
            logger.error(f"🔴 Failed to initialize templates directory: {e}")
            logger.error(f"   Full exception: {traceback.format_exc()}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Templates directory initialization failed: {str(e)}"
            )
    
    def validate_input_data(self, opportunity_data: Dict[str, Any], sections: list, funder_notes: Optional[str] = None) -> Dict[str, Any]:
        """Validate input data before processing"""
        validation_result = {
            'valid': True,
            'warnings': [],
            'errors': []
        }
        
        # Validate opportunity data
        if not isinstance(opportunity_data, dict):
            validation_result['errors'].append("Opportunity data must be a dictionary")
            validation_result['valid'] = False
            return validation_result
        
        # Check required opportunity fields
        required_fields = ['title']
        for field in required_fields:
            if not opportunity_data.get(field):
                validation_result['warnings'].append(f"Missing or empty opportunity field: {field}")
        
        # Validate sections
        if not isinstance(sections, list):
            validation_result['errors'].append("Sections must be a list")
            validation_result['valid'] = False
        elif not sections:
            validation_result['errors'].append("At least one section is required")
            validation_result['valid'] = False
        else:
            for i, section in enumerate(sections):
                if not isinstance(section, dict):
                    validation_result['errors'].append(f"Section {i+1} must be a dictionary")
                    validation_result['valid'] = False
                    continue
                
                if not section.get('heading', '').strip():
                    validation_result['errors'].append(f"Section {i+1} missing heading")
                    validation_result['valid'] = False
                
                if not section.get('instruction', '').strip():
                    validation_result['warnings'].append(f"Section {i+1} missing instruction")
                
                # Check for excessively long content
                heading_length = len(section.get('heading', ''))
                instruction_length = len(section.get('instruction', ''))
                
                if heading_length > 200:
                    validation_result['warnings'].append(f"Section {i+1} heading very long: {heading_length} chars")
                if instruction_length > 2000:
                    validation_result['warnings'].append(f"Section {i+1} instruction very long: {instruction_length} chars")
        
        # Validate funder notes length
        if funder_notes and len(funder_notes) > 5000:
            validation_result['warnings'].append(f"Funder notes very long: {len(funder_notes)} chars")
        
        # Log validation results
        if validation_result['warnings']:
            logger.warning(f"⚠️ Input validation warnings: {'; '.join(validation_result['warnings'])}")
        if validation_result['errors']:
            logger.error(f"🔴 Input validation errors: {'; '.join(validation_result['errors'])}")
        
        return validation_result
    
    def create_proposal_template(
        self, 
        opportunity_data: Dict[str, Any],
        sections: list,
        funder_notes: Optional[str] = None
    ) -> str:
        """
        Create a .docx proposal template with opportunity context and QA-defined sections
        
        Args:
            opportunity_data: Funding opportunity data for context
            sections: List of sections with heading and instruction
            funder_notes: Optional funder-specific notes
            
        Returns:
            str: Filename of the generated document
        """
        template_id = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:20]
        filename = None
        filepath = None
        
        try:
            logger.info(f"📄 [{template_id}] Starting proposal template generation")
            logger.info(f"   Opportunity title: {opportunity_data.get('title', 'Unknown')}")
            logger.info(f"   Number of sections: {len(sections)}")
            logger.info(f"   Funder notes: {'Yes' if funder_notes else 'No'}")
            
            # Validate input data
            validation_result = self.validate_input_data(opportunity_data, sections, funder_notes)
            if not validation_result['valid']:
                logger.error(f"🔴 [{template_id}] Input validation failed:")
                for error in validation_result['errors']:
                    logger.error(f"   - {error}")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Input validation failed: {'; '.join(validation_result['errors'])}"
                )
            
            # Ensure directory exists and is writable
            self._ensure_templates_directory()
            
            logger.info(f"📝 [{template_id}] Creating Word document...")
            
            # Create new document
            try:
                doc = Document()
                logger.info(f"✅ [{template_id}] Word document created successfully")
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to create Word document: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to create Word document: {str(e)}"
                )
            
            # Add title
            title = opportunity_data.get('title', 'Funding Opportunity Proposal')
            logger.info(f"📋 [{template_id}] Adding title: {title}")
            try:
                title_paragraph = doc.add_heading(f"Proposal Template: {title}", 0)
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to add title: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to add document title: {str(e)}"
                )
            
            # Add opportunity metadata section
            logger.info(f"📊 [{template_id}] Adding opportunity information section...")
            try:
                doc.add_heading('Opportunity Information', level=1)
                
                # Create opportunity details table
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                
                # Add opportunity details
                metadata_fields = [
                    ('Donor/Funder', opportunity_data.get('donor', 'Not specified')),
                    ('Deadline', opportunity_data.get('deadline', 'Not specified')),
                    ('Funding Amount', opportunity_data.get('amount', 'Not specified')),
                    ('Location/Eligibility', opportunity_data.get('location', 'Not specified')),
                    ('Themes/Focus Areas', ', '.join(opportunity_data.get('themes', [])) if isinstance(opportunity_data.get('themes'), list) else opportunity_data.get('themes', 'Not specified')),
                    ('Opportunity URL', opportunity_data.get('opportunity_url', 'Not provided'))
                ]
                
                for label, value in metadata_fields:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[0].paragraphs[0].runs[0].bold = True
                    row_cells[1].text = str(value)
                
                logger.info(f"✅ [{template_id}] Added {len(metadata_fields)} metadata fields")
                
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to add opportunity information: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to add opportunity information: {str(e)}"
                )
            
            # Add some spacing
            doc.add_paragraph('')
            
            # Add funder notes if provided
            if funder_notes and funder_notes.strip():
                logger.info(f"📝 [{template_id}] Adding funder notes ({len(funder_notes)} chars)")
                try:
                    doc.add_heading('Funder Requirements & Notes', level=1)
                    funder_para = doc.add_paragraph(funder_notes.strip())
                    funder_para.style = 'Intense Quote'
                    doc.add_paragraph('')
                    logger.info(f"✅ [{template_id}] Added funder notes")
                except Exception as e:
                    logger.error(f"🔴 [{template_id}] Failed to add funder notes: {e}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"Failed to add funder notes: {str(e)}"
                    )
            
            # Add instruction paragraph
            logger.info(f"📋 [{template_id}] Adding instructions section...")
            try:
                instruction_para = doc.add_paragraph()
                instruction_para.add_run('Instructions: ').bold = True
                instruction_para.add_run('This template provides the structure for your proposal. Replace the instructional text below with your actual content. Each section includes guidance on what to include.')
                doc.add_paragraph('')
                logger.info(f"✅ [{template_id}] Added instructions")
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to add instructions: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to add instructions: {str(e)}"
                )
            
            # Add user-defined sections
            logger.info(f"📄 [{template_id}] Adding {len(sections)} proposal sections...")
            try:
                doc.add_heading('Proposal Sections', level=1)
                
                for i, section in enumerate(sections, 1):
                    section_heading = section.get('heading', f'Section {i}')
                    section_instruction = section.get('instruction', 'No instruction provided')
                    
                    logger.info(f"   📝 [{template_id}] Adding section {i}: {section_heading}")
                    
                    # Add section heading
                    section_heading_para = doc.add_heading(f"{i}. {section_heading}", level=2)
                    
                    # Add instruction as placeholder text
                    instruction_para = doc.add_paragraph()
                    instruction_para.add_run('[INSTRUCTION] ').bold = True
                    instruction_para.add_run(section_instruction)
                    instruction_para.style = 'Subtle Emphasis'
                    
                    # Add placeholder for actual content
                    content_para = doc.add_paragraph()
                    content_para.add_run('[Your content for this section goes here]')
                    content_para.style = 'Quote'
                    
                    # Add spacing between sections
                    doc.add_paragraph('')
                
                logger.info(f"✅ [{template_id}] Added all {len(sections)} sections")
                
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to add proposal sections: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to add proposal sections: {str(e)}"
                )
            
            # Add footer with generation info
            logger.info(f"📋 [{template_id}] Adding footer information...")
            try:
                doc.add_page_break()
                footer_heading = doc.add_heading('Template Information', level=2)
                footer_para = doc.add_paragraph()
                footer_para.add_run('Generated: ').bold = True
                footer_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                footer_para.add_run('\nBy: ').bold = True
                footer_para.add_run('ReqAgent Proposal Template Generator')
                footer_para.add_run('\nOpportunity Source: ').bold = True
                footer_para.add_run(opportunity_data.get('opportunity_url', 'Not provided'))
                logger.info(f"✅ [{template_id}] Added footer information")
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to add footer: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to add footer information: {str(e)}"
                )
            
            # Generate unique filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()[:50]
            filename = f"proposal_template_{safe_title}_{timestamp}.docx"
            filepath = os.path.join(self.templates_dir, filename)
            
            logger.info(f"💾 [{template_id}] Saving document to: {filepath}")
            
            # Save document
            try:
                doc.save(filepath)
                logger.info(f"✅ [{template_id}] Document saved successfully")
            except PermissionError as e:
                logger.error(f"🔴 [{template_id}] Permission denied when saving document: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Permission denied when saving template. Check file system permissions: {str(e)}"
                )
            except OSError as e:
                logger.error(f"🔴 [{template_id}] OS error when saving document: {e}")
                logger.error(f"   File path: {filepath}")
                logger.error(f"   Directory exists: {os.path.exists(self.templates_dir)}")
                logger.error(f"   Directory writable: {os.access(self.templates_dir, os.W_OK)}")
                
                # Check disk space again
                try:
                    disk_usage = shutil.disk_usage(self.templates_dir)
                    free_space_mb = disk_usage.free / (1024 * 1024)
                    logger.error(f"   Available disk space: {free_space_mb:.1f}MB")
                except:
                    logger.error(f"   Could not check disk space")
                
                raise HTTPException(
                    status_code=status.HTTP_507_INSUFFICIENT_STORAGE,
                    detail=f"Failed to save template due to file system error: {str(e)}"
                )
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Unexpected error when saving document: {e}")
                logger.error(f"   Full exception: {traceback.format_exc()}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Unexpected error when saving template: {str(e)}"
                )
            
            # Verify file was created and get file size
            try:
                if not os.path.exists(filepath):
                    raise FileNotFoundError(f"Template file was not created: {filepath}")
                
                file_size = os.path.getsize(filepath)
                logger.info(f"📏 [{template_id}] Template file size: {file_size:,} bytes")
                
                if file_size < 1000:  # Less than 1KB seems suspicious
                    logger.warning(f"⚠️ [{template_id}] Template file seems unusually small: {file_size} bytes")
                
            except Exception as e:
                logger.error(f"🔴 [{template_id}] Failed to verify saved file: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Template was saved but verification failed: {str(e)}"
                )
            
            logger.info(f"✅ [{template_id}] Generated proposal template: {filename}")
            return filename
            
        except HTTPException:
            # Re-raise HTTP exceptions as they already have proper error details
            raise
        except Exception as e:
            logger.error(f"🔴 [{template_id}] Unexpected error in template generation: {e}")
            logger.error(f"   Full exception: {traceback.format_exc()}")
            logger.error(f"   Template generation context:")
            logger.error(f"     Opportunity data keys: {list(opportunity_data.keys()) if isinstance(opportunity_data, dict) else 'Not a dict'}")
            logger.error(f"     Sections count: {len(sections) if isinstance(sections, list) else 'Not a list'}")
            logger.error(f"     Funder notes length: {len(funder_notes) if funder_notes else 0}")
            logger.error(f"     Templates directory: {self.templates_dir}")
            logger.error(f"     Target filename: {filename}")
            logger.error(f"     Target filepath: {filepath}")
            
            # Clean up partial file if it exists
            if filepath and os.path.exists(filepath):
                try:
                    os.remove(filepath)
                    logger.info(f"🧹 [{template_id}] Cleaned up partial file: {filepath}")
                except:
                    logger.warning(f"⚠️ [{template_id}] Could not clean up partial file: {filepath}")
            
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate proposal template: {str(e)}"
            )

@router.get("/proposal-template/start", response_class=HTMLResponse)
async def proposal_template_start(
    request: Request,
    record_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: str = Depends(require_admin_auth)
):
    """
    Display the proposal template creation form
    """
    try:
        opportunity_data = None
        
        # If record_id is provided, fetch the opportunity data
        if record_id:
            opportunity = db.query(FundingOpportunity).filter(
                FundingOpportunity.id == record_id
            ).first()
            
            if not opportunity:
                raise HTTPException(
                    status_code=404,
                    detail=f"Funding opportunity with ID {record_id} not found"
                )
            
            # Check if opportunity is approved
            if opportunity.status != StatusEnum.approved:
                raise HTTPException(
                    status_code=400,
                    detail=f"Opportunity must be approved before creating proposal template. Current status: {opportunity.status.value}"
                )
            
            opportunity_data = {
                "id": opportunity.id,
                "title": opportunity.json_data.get('title', 'Unknown') if opportunity.json_data else 'Unknown',
                "donor": opportunity.json_data.get('donor', 'Unknown') if opportunity.json_data else 'Unknown',
                "deadline": opportunity.json_data.get('deadline', 'Unknown') if opportunity.json_data else 'Unknown',
                "amount": opportunity.json_data.get('amount', 'Unknown') if opportunity.json_data else 'Unknown',
                "themes": opportunity.json_data.get('themes', []) if opportunity.json_data else [],
                "location": opportunity.json_data.get('location', 'Unknown') if opportunity.json_data else 'Unknown',
                "opportunity_url": opportunity.json_data.get('opportunity_url', opportunity.source_url) if opportunity.json_data else opportunity.source_url
            }
        
        # Get all approved opportunities for selection
        approved_opportunities = db.query(FundingOpportunity).filter(
            FundingOpportunity.status == StatusEnum.approved
        ).order_by(FundingOpportunity.created_at.desc()).limit(50).all()
        
        opportunities_list = []
        for opp in approved_opportunities:
            opp_data = {
                "id": opp.id,
                "title": opp.json_data.get('title', 'Unknown') if opp.json_data else 'Unknown',
                "donor": opp.json_data.get('donor', 'Unknown') if opp.json_data else 'Unknown',
                "created_at": opp.created_at.strftime("%Y-%m-%d")
            }
            opportunities_list.append(opp_data)
        
        return templates.TemplateResponse(
            "proposal_template_form.html",
            {
                "request": request,
                "opportunity": opportunity_data,
                "opportunities": opportunities_list,
                "current_user": current_user,
                "page_title": "Create Proposal Template"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error in proposal template start: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to load proposal template form: {str(e)}"
        )

@router.post("/proposal-template/generate", response_model=ProposalTemplateResponse)
async def generate_proposal_template(
    request: CreateProposalTemplateRequest,
    db: Session = Depends(get_db),
    current_user: str = Depends(require_admin_auth)
):
    """
    Generate a .docx proposal template based on QA-defined sections
    """
    try:
        logger.info(f"🚀 Generating proposal template for record ID: {request.record_id} (User: {current_user})")
        
        # Fetch the funding opportunity
        opportunity = db.query(FundingOpportunity).filter(
            FundingOpportunity.id == request.record_id
        ).first()
        
        if not opportunity:
            raise HTTPException(
                status_code=404,
                detail=f"Funding opportunity with ID {request.record_id} not found"
            )
        
        # Check if opportunity is approved
        if opportunity.status != StatusEnum.approved:
            raise HTTPException(
                status_code=400,
                detail=f"Opportunity must be approved before creating proposal template. Current status: {opportunity.status.value}"
            )
        
        # Extract opportunity data
        opportunity_data = opportunity.json_data or {}
        opportunity_data['opportunity_url'] = opportunity_data.get('opportunity_url', opportunity.source_url)
        
        # Validate sections
        if not request.sections:
            raise HTTPException(
                status_code=400,
                detail="At least one section is required to generate a proposal template"
            )
        
        # Convert sections to dict format
        sections_data = [
            {
                "heading": section.heading,
                "instruction": section.instruction
            }
            for section in request.sections
        ]
        
        # Generate the template
        generator = ProposalTemplateGenerator()
        filename = generator.create_proposal_template(
            opportunity_data=opportunity_data,
            sections=sections_data,
            funder_notes=request.funder_notes
        )
        
        # Create download URL
        download_url = f"/admin/proposal-template/download/{filename}"
        
        return ProposalTemplateResponse(
            success=True,
            message=f"Successfully generated proposal template with {len(sections_data)} sections",
            filename=filename,
            download_url=download_url,
            timestamp=datetime.now().isoformat(),
            opportunity_title=opportunity_data.get('title', 'Unknown Opportunity')
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error generating proposal template: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate proposal template: {str(e)}"
        )

@router.get("/proposal-template/download/{filename}")
async def download_proposal_template(
    filename: str,
    current_user: str = Depends(require_admin_auth)
):
    """
    Download a generated proposal template
    """
    try:
        file_path = os.path.join(tempfile.gettempdir(), "reqagent_templates", filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=404,
                detail="Template file not found"
            )
        
        logger.info(f"📥 Downloading proposal template: {filename} (User: {current_user})")
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error downloading template: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to download template: {str(e)}"
        ) 